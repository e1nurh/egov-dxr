import requests
from bs4 import BeautifulSoup as bs
from difflib import SequenceMatcher
import time
import Levenshtein

class Scraper():
	def __init__(self):
		self.egov_services_url = "https://www.e-gov.az/az/services/"
		self.dxr_services_url = "http://dxr.az/dovlet-xidmetleri?page={}"
		self.egov_services = []
		self.dxr_services = []
		self.egov_dxr = []
		self.not_in_egov = []
		self.not_in_dxr = []
		self.egov_dxr_with_percent = []

	def egov(self):
		a = time.time()
		r = requests.get(self.egov_services_url, verify=False)
		b = bs(r.content, 'html.parser')
		services = b.find('ul', {"class":"menu-accordion"})
		links = services.find_all('a', text=True)
		for link in links:
			self.egov_services.append(link.getText())
		print("Avarage time for e-gov.az scraping", time.time() - a)


	def dxr(self):
		a = time.time()
		for n in range(1, 50):
			print(self.dxr_services_url.format(str(n)))
			r = requests.get(self.dxr_services_url.format(str(n)))
			b = bs(r.content, 'html.parser')
			check_page = b.find('h3', {"class":"center"})
			if check_page:
				break
			services_ul = b.find("ul", {"id":"search_list"})
			services_texts = services_ul.find_all("a", {"class":"name"})
			for i in services_texts:
				clear_sentence = i.getText().replace("  ", "").replace("\t", "").replace("\n", "") 
				if clear_sentence[0] == " ":
					clear_sentence = clear_sentence[1:]
				if clear_sentence[-1] == " ":
					clear_sentence = clear_sentence[:-1]
				self.dxr_services.append(clear_sentence)
		print("Avarage time for dxr.az scraping:", time.time() - a)

	def get_intersection(self):
		a = time.time()
		for egov_service in self.egov_services:
			for dxr_service in self.dxr_services:
				match_percent = round(Levenshtein.ratio(egov_service, dxr_service), 2)*100
				if match_percent>95:
					if egov_service not in self.egov_dxr_with_percent:
						self.egov_dxr_with_percent.append(egov_service) 
						break
		print("Avarage time for calculating intersections", time.time() - a)

	def get_intercetion_without_matching(self):
		for egov_service in self.egov_services:
			if egov_service in self.dxr_services:
				self.egov_dxr.append(egov_service)

	def write_to_docx(self):
		print("scraping e-gov.az ...")
		self.egov()
		print("scraping dxr.az ...")
		self.dxr()
		print("Calculating intersections of dxr.az and e-gov.az")
		self.get_intersection()
		from docx import Document
		document = Document()
		document.add_heading('dxr.az və e-gov.az saytlarında olan ortaq xidmətlərin siyahısı', 0)
		p = document.add_paragraph('Ortaq xidmətlərin sayı - {}'.format(str(len(self.egov_dxr_with_percent))))
		for i in self.egov_dxr_with_percent:
			document.add_paragraph(
		    	i, style='List Number'
			)
		document.save("intersection.docx")

		self.get_intercetion_without_matching()
		document = Document()
		document.add_heading("dxr.az və e-gov.az saytlarında olan ortaq xidmətlərin siyahısı (\
								analiz etmədən)", 0)
		document.add_paragraph("Ortaq xidmətlərin siyahısı - {}".format(str(len(self.egov_dxr))))
		for x in self.egov_dxr:
			document.add_paragraph(x, style='List Number')
		document.save("intercetion_without_analyze.docx")

	def all_services(self):
		a = time.time()
		self.egov_dxr = list(set(self.egov_services+self.dxr_services))
		print(time.time() - a)

	def not_in_dxr_func(self):
		for egov_service in self.egov_services:
			if egov_service not in self.dxr_services:
				self.not_in_dxr.append(egov_service)

	def not_in_egov_func(self):
		for dxr_service in self.dxr_services:
			if dxr_service not in self.egov_services:
				self.not_in_egov.append(dxr_service)


